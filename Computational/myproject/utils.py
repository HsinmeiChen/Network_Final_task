import base64
import configparser
import re
import openai
from myproject.assistant_api import *
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn  # Import the qn function for namespace handling

config = configparser.ConfigParser()
config.read('config.ini')


def Struct_Step_Question_Attr(content_list):
    """
    Combines content lists into structured sub-objects:
    - Guides are stored as an array of {Description, Answers}.
    - Questions include Description, Hint, Images, and a single Answer.
    - Debug includes a single Description.
    """
    structured_content = {"Guides": {}}  # Initialize Guides as a dictionary
    for section in content_list:
        section_name = section.get("name", "Other").split('-')[0]  # Main section name (e.g., Guide1, Debug)
        section_part = section.get("name", "Other").split('-')[1:]  # Subsection name (e.g., Description, Ans)

        # If "Question", ensure single structured content
        if section_name == "Question":
            if "Question" not in structured_content:
                structured_content["Question"] = {"Description": "", "Hint": "", "Images": [], "Answer": ""}
            part_key = "Answer" if "Answer" in section_part else "Hint" if "Hint" in section_part else "Description"
            for content in section["value"]:
                if content["type"] == "text":
                    structured_content["Question"][part_key] += content["value"]
                elif content["type"] == "image":
                    structured_content["Question"]["Images"].append(content["value"])
        
        # If "Guide", store as an array with Description and Answers
        elif section_name.startswith("Guide"):
            if section_name not in structured_content["Guides"]:
                structured_content["Guides"][section_name] = {"Description": "", "Answers": [], "Images": []}
            for content in section["value"]:
                if content["type"] == "text":
                    if "Answer" in section_part:
                        structured_content["Guides"][section_name]["Answers"].append(content["value"].strip())
                    else:
                        structured_content["Guides"][section_name]["Description"] += content["value"]
                elif content["type"] == "image":
                    structured_content["Guides"][section_name]["Images"].append(content["value"])

        # If "Debug", ensure single structured content
        elif section_name == "Debug":
            if "Debug" not in structured_content:
                structured_content["Debug"] = {"Description": ""}
            for content in section["value"]:
                if content["type"] == "text":
                    structured_content["Debug"]["Description"] += content["value"]

    return structured_content

def load_teaching_plan(docx_path):
    doc = Document(docx_path)
    teaching_plan = []
    step_data = None  # Initialize the current step

    # Regular expressions
    section_pattern = re.compile(r'【(.*?)】')  # Match sections like 【Question-Description】
    step_pattern = re.compile(r'^Step\s*(\d+)')  # Match "Step1", "Step2", etc.
    current_section = None
    statement_buffer = []  # Buffer to store statement content

    def get_image_dimensions(element):
        """Extract image dimensions from Word drawing element"""
        # Find inline or anchor drawing properties
        drawing_prop = element.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
        if drawing_prop is not None:
            # Convert EMUs (English Metric Units) to pixels
            # 1 EMU = 1/914400 inch, assuming 96 DPI
            emu_to_pixels = lambda emu: int(int(emu) * 96 / 914400)
            width = emu_to_pixels(drawing_prop.get('cx', '0'))
            height = emu_to_pixels(drawing_prop.get('cy', '0'))
            return width, height
        return None, None

    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text.replace("\n", "<br>").replace("\r", "<br>").strip()

        # Check for a new Step
        step_match = step_pattern.match(paragraph_text)
        if step_match:
            # Finalize the previous step
            if step_data:
                step_data["statement"] = " ".join(statement_buffer).strip()
                step_data["content"] = Struct_Step_Question_Attr(step_data["content"])
                teaching_plan.append(step_data)

            # Start a new step
            step_data = {"step": paragraph_text, "statement": "", "content": []}
            statement_buffer = []  # Reset statement buffer
            current_section = None
            continue

        # Match sections like 【Question-Description】
        section_match = section_pattern.match(paragraph_text)
        if section_match:
            current_section = section_match.group(1)
            step_data["content"].append({"name": current_section, "value": []})
        elif current_section:
            # Add content to the current section
            step_data["content"][-1]["value"].append({"type": "text", "value": paragraph_text + "<br>"})
        else:
            # Append to the statement buffer if outside any 【...】 section
            statement_buffer.append(paragraph_text)

        # Process images in runs
        for run in paragraph.runs:
            for drawing_elem in run._element.findall('.//{}drawing'.format(qn('w:'))):
                blip = drawing_elem.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                if blip is not None:
                    image_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if image_id:
                        # Get image dimensions
                        width, height = get_image_dimensions(drawing_elem)
                        
                        # Get image data
                        image_part = doc.part.related_parts[image_id]
                        image_data = image_part.blob
                        image_base64 = base64.b64encode(image_data).decode('utf-8')
                        
                        # Create image HTML with dimensions if available
                        img_html = '<img src="data:image/png;base64,{}"'.format(image_base64)
                        if width and height:
                            img_html += ' width="{}" height="{}"'.format(width, height)
                        img_html += ' alt="Embedded Image">'
                        
                        if current_section:
                            step_data["content"][-1]["value"].append(
                                {"type": "image", "value": img_html}
                            )

    # Finalize the last step
    if step_data:
        step_data["statement"] = " ".join(statement_buffer).strip()
        step_data["content"] = Struct_Step_Question_Attr(step_data["content"])
        teaching_plan.append(step_data)

    return teaching_plan


# 將教案解析為佇列，並處理特殊指令
def initialize_teaching_plan_queue(teaching_plan):
    """
    Converts the teaching_plan into a queue structure.
    Each step in the queue contains 'step', 'statement', and 'content'.
    """
    step_queue = []  # Initialize the queue to store steps

    for step_data in teaching_plan:
        # Each step_data should already have 'step', 'statement', and 'content'
        step_entry = {
            "step": step_data["step"],
            "statement": step_data["statement"],
            "content": step_data["content"]
        }
        step_queue.append(step_entry)

    return step_queue


# 處理教案中的特殊指令
def process_special_instructions(content):
    step_text = ""  # 用於串接文字段落和圖片

    for content_item in content:
        if content_item['type'] == 'text':
            # 將文字串接成完整段落
            content = content_item['value'].replace("\n", "<br>")
            step_text += content_item['value']
        elif content_item['type'] == 'image':
            # 若有圖片，將圖片轉換為 <img> 標籤格式並直接串接
            image_html = f'<img src="{content_item["value"]}" alt="Step Image">'
            step_text += image_html + " "

    # 去除多餘空白並返回完整的字串
    return step_text.strip()

# 呼叫 OpenAI ChatGPT API
def ask_chatgptAPI(company_name,inputJson,_serviceName,model):
    rsbody = Service_Assistant.api_Assistant_OneShot(
        company_name=company_name,
        config=config,
        serviceName=_serviceName,
        model=model,
        message=inputJson
    )
    return rsbody

def ask_chatgpt(company_name,inputMsg,_assistantName,assistantInstruction,model):
    rsbody = Service_Assistant.Assistant_OneShot(
        company_name=company_name,
        config=config,
        assistantname=_assistantName,
        assistant_instruction=assistantInstruction,
        model=model,
        message=inputMsg
    )
    return rsbody

# 提取並解析使用者的名字
def extract_name(user_message):
    possible_name = user_message.strip()
    if possible_name.lower().startswith("我叫") or possible_name.lower().startswith("我是"):
        possible_name = possible_name[2:].strip()
    elif possible_name.lower().startswith("叫我"):
        possible_name = possible_name[3:].strip()
    return possible_name

# 檢查使用者回應中是否包含否定句
def contains_negation(user_message):
    negations = ["看完了", "不知道", "我不懂", "我不清楚", "no", "沒有", "No", "我不知道", "我不會", "不清楚", "不了解", "不明白", "不確定", "沒聽過"]
    for negation in negations:
        if negation in user_message:
            return True
    return False

def convert_hard_breaks_to_soft_and_save(docx_path):
    # 開啟原始 docx 文件
    doc = Document(docx_path)
    content_with_soft_breaks = ""

    # 遍歷每個段落，將硬斷行替換為 <br>
    for i, paragraph in enumerate(doc.paragraphs):
        paragraph_text = paragraph.text.strip()
        
        # 累加段落文字並在段落之間加上 <br>（最後一段落不需要）
        if paragraph_text:
            content_with_soft_breaks += paragraph_text
            if i < len(doc.paragraphs) - 1:
                content_with_soft_breaks += "<br>"

    # 創建新的 docx 文件
    new_doc = Document()
    new_doc.add_paragraph(content_with_soft_breaks)

    # 將轉換結果儲存回原檔案
    new_doc.save(docx_path)

def read_docx_content(docx_path):
    """
    讀取指定 .docx 文件的所有內容，並回傳為字串。
    
    Args:
        docx_path (str): .docx 文件的路徑。
    
    Returns:
        str: 文件的所有內容，段落以換行符分隔。
    """
    # 開啟指定 .docx 文件
    doc = Document(docx_path)
    content = []

    # 遍歷所有段落並將內容加入清單
    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text.strip()
        if paragraph_text:
            content.append(paragraph_text)
    
    # 將所有段落連接為單一字串，並以換行符分隔
    return "\n".join(content)